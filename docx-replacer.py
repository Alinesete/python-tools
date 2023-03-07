from docx import Document
import os
import datetime
import calendar
import shutil


def main():
    template_file_path = './Modelos-pol'
    output_file_path = ('./Politicas-Procedimentos')

    default_num_cart = "NUM_DOC"
    numero_cart = input("Qual o número do cartório: ")
    default_raz_soc= "CONTROLADOR-CLIENTE"
    raz_soc = input("Digite a razão social do cliente: ")
    default_fantasia = "CONTROLADOR-FANTASIA"
    fantasia = input("Digite o nome-fantasia do cliente: ")
    default_cidade = "COD_MUNICIPIO"
    cidade = input("Digite a cidade: ")
    default_uf = "COD_UF"
    uf = "GO"
    data_default = "dd/mm/aaaa"
    data_hoje = data_atual.strftime("%d/%m/%Y")

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == '__main__':
    main()
