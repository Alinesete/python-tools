import os
import docx
import datetime
import calendar
import shutil
import locale
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

data_atual = datetime.datetime.now()

locale.setlocale(locale.LC_ALL, 'pt_PT.utf8')
nome_mes = calendar.month_name[data_atual.month].upper()

diretorio = "./Base-docs"

# Dicionário de substituições
default_num_cart = "NUM_DOC"
default_cidade = "CONTROLADOR-MUNICIPIO"
default_comarca = "CONTROLADOR-COMARCA"
default_raz_soc= "CONTROLADOR-CLIENTE"
default_email = "CONTROLADOR-EMAIL"
default_uf = "CONTROLADOR-ESTADO"
default_site = "CONTROLADOR-SITE"
data_default = "dd/mm/aaaa"
data_extenso_default = "DD-DE-MM-DE-AAAA"
mes_default = "MES/ANO"

# Inputs
numero_cart = input("Qual o número do cartório: ")
raz_soc = input("Digite a razão social do cliente: ")
cidade = input("Digite a cidade ou município: ")
comarca = input("Digite a cormaca (Caso não haja, deixar em branco): ")
uf = input("Digite a UF: ").upper()
email = input("Digite o email (Caso não haja, deixar em branco): ").lower()
site = input("Digite o site (Caso não haja, deixar em branco): ").lower()

# Data atual
data_hoje = data_atual.strftime("%d/%m/%Y")
data_extenso_hoje = data_atual.strftime('%d de %B de %Y')
mes_hoje = f"{nome_mes}/{str(data_atual.year)[-2:]}"

# Verificação de input em branco
if not comarca:
    comarca = cidade
if not email:
    email = "cartorio2oficio@email.com"
if not site:
    site = "«CLIENTE_DOMINIO_SITE»"


# Cria a nova pasta pegando os modelos e adicionando o número do cartório na frente
diretorio_novo = "./Politicas-Procedimentos_" + numero_cart
if not os.path.exists(diretorio_novo):
    os.makedirs(diretorio_novo)

# Todo o processo de substituir
def main():

    # Percorrer todos os arquivos no modelo e copiar no diretório novo
    for arquivo in os.listdir(diretorio):

        # Copia todos os arquivos terminados em .docx para a pasta nova e os percorre
        if arquivo.endswith(".docx"):
            caminho_arquivo = os.path.join(diretorio, arquivo)

            caminho_arquivo_novo = os.path.join(diretorio_novo, numero_cart + " - " + arquivo)
            shutil.copy(caminho_arquivo, caminho_arquivo_novo)

            doc = docx.Document(caminho_arquivo_novo)

            head = doc.sections[0].header

            # Dicionário de substituições default : novo
            Dictionary = {default_num_cart : numero_cart, default_raz_soc : raz_soc, default_cidade : cidade, default_uf : uf, default_site : site, default_email : email, mes_default : mes_hoje, data_extenso_default : data_extenso_hoje, default_comarca : comarca, data_default : data_hoje}

            # Substituição
            for i in Dictionary:

                # Percorre todos os parágrafos normais
                for p in doc.paragraphs:
                    for run in p.runs:
                        if i in run.text:
                            text = run.text.replace(i, Dictionary[i])
                            run.text = ""
                            for idx, part in enumerate(text.split(i)):
                                run.add_text(part)
                                if idx < len(text.split(i)) - 1:
                                    run.add_text(i)
                                    font = run.font
                                    font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                                    font.highlight_color = WD_COLOR_INDEX.YELLOW
                                    font.bold = True
                        
                # Percorre todas as tabelas no corpo
                for tabela in doc.tables:
                    for linha in tabela.rows:
                        for celula in linha.cells:
                            for i in Dictionary.keys():

                                # Mudança apenas no dd/mm/aaaa, usa a formatação do estilo "Normal"
                                if data_default in celula.text:
                                    celula.text = celula.text.replace(data_default, data_hoje)
                                
                                # Mudança com base no dicionário sem perder formatação
                                for key in Dictionary.keys():
                                    for paragraph in celula.paragraphs:
                                        for run in paragraph.runs:
                                            if key in run.text:
                                                run.text = run.text.replace(key, Dictionary[key])

                #Percorre todas as tabelas no cabeçalho
                for tabela in head.tables:
                    for linha in tabela.rows:
                        for celula in linha.cells:

                            # Mudança com base no dicionário sem perder formatação
                            for i in Dictionary.keys():
                                for paragraph in celula.paragraphs:
                                    for run in paragraph.runs:
                                        if i in run.text:
                                            run.text = run.text.replace(i, Dictionary[i])

            doc.save(caminho_arquivo_novo)

if __name__ == '__main__':
    main()